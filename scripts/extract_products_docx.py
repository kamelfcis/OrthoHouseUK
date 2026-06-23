#!/usr/bin/env python3
"""Extract products and images from Docs/products.docx"""
import json
import re
import zipfile
import os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
DOC_PATH = ROOT / "Docs" / "products.docx"
OUT_DIR = ROOT / "scripts" / "tmp" / "product_images"
OUT_JSON = ROOT / "scripts" / "tmp" / "extracted_products.json"


def extract_images(doc_path: Path, out_dir: Path) -> dict:
    """Extract embedded images from docx zip, return rId -> filename map."""
    out_dir.mkdir(parents=True, exist_ok=True)
    image_map = {}
    with zipfile.ZipFile(doc_path) as z:
        media_files = sorted(n for n in z.namelist() if n.startswith("word/media/"))
        for i, media_path in enumerate(media_files):
            ext = Path(media_path).suffix or ".png"
            filename = f"image_{i+1:03d}{ext}"
            dest = out_dir / filename
            dest.write_bytes(z.read(media_path))
            image_map[media_path] = filename
    return image_map


def get_paragraph_image_rids(paragraph):
    """Find image relationship ids in a paragraph."""
    rids = []
    for run in paragraph.runs:
        blips = run._element.findall(".//" + qn("a:blip"))
        for blip in blips:
            rid = blip.get(qn("r:embed"))
            if rid:
                rids.append(rid)
    return rids


def build_rid_to_media(doc: Document, doc_path: Path) -> dict:
    """Map relationship id to media zip path."""
    rid_map = {}
    part = doc.part
    for rel in part.rels.values():
        if "image" in rel.reltype:
            target = rel.target_ref
            if not target.startswith("word/"):
                target = "word/" + target.lstrip("/")
            rid_map[rel.rId] = target
    return rid_map


def is_heading(text: str) -> bool:
    """Heuristic: short line without trailing colon, not a spec bullet."""
    t = text.strip()
    if not t or len(t) > 120:
        return False
    if t.endswith(":"):
        return False
    if t[0].islower():
        return False
    # spec-like patterns
    if re.match(r"^(Fixed|Low|Anatomic|Pure|Titanium|Self|Torx|Variable|Left|Rounded|Special)", t):
        return False
    return True


def parse_products(doc: Document, rid_to_media: dict, media_filename_map: dict) -> list:
    products = []
    current = None
    current_category = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Category headers (often short, title case, no colon)
        style = para.style.name if para.style else ""
        if style.startswith("Heading") or (len(text) < 40 and text.isupper() is False and ":" not in text and para.runs and para.runs[0].bold):
            pass  # fall through to heuristics

        # Get images in this paragraph
        rids = get_paragraph_image_rids(para)
        para_images = []
        for rid in rids:
            media_path = rid_to_media.get(rid)
            if media_path and media_path in media_filename_map:
                para_images.append(media_filename_map[media_path])

        # Product name heuristic: ends with colon OR is a short title after category
        if text.endswith(":"):
            name = text[:-1].strip()
            if current:
                products.append(current)
            current = {
                "name": name,
                "category": current_category,
                "description_lines": [],
                "specifications": [],
                "images": para_images[:],
            }
            continue

        # Possible category (single word/title before products)
        if is_heading(text) and not current and not text.endswith("."):
            # Could be brand/category like "Astrolabe"
            if len(text.split()) <= 3:
                current_category = text
                continue

        # New product without colon - short standalone title
        if is_heading(text) and current is None:
            current_category = None
            current = {
                "name": text,
                "category": current_category,
                "description_lines": [],
                "specifications": [],
                "images": para_images[:],
            }
            continue

        if current is None:
            # First line might be category
            if len(text.split()) <= 2:
                current_category = text
            continue

        # Append images to current product
        if para_images:
            current["images"].extend(para_images)

        # Classify as spec (semicolon ending or bullet-like) vs description
        if text.endswith(";") or (len(text) < 100 and ";" in text):
            current["specifications"].append(text.rstrip(";"))
        else:
            current["description_lines"].append(text)

    if current:
        products.append(current)

    # Post-process
    for p in products:
        p["description"] = " ".join(p["description_lines"]).strip()
        p["specifications_text"] = "\n".join(p["specifications"]).strip()
        p["product_code"] = re.sub(r"[^A-Za-z0-9]+", "-", p["name"]).strip("-").upper()[:50]
        p["images"] = list(dict.fromkeys(p["images"]))  # dedupe preserve order
        del p["description_lines"]

    return products


def main():
    doc = Document(DOC_PATH)
    media_filename_map = extract_images(DOC_PATH, OUT_DIR)
    rid_to_media = build_rid_to_media(doc, DOC_PATH)

    products = parse_products(doc, rid_to_media, media_filename_map)

    # Also dump raw paragraphs for debugging
    raw_paragraphs = [
        {"i": i, "style": p.style.name if p.style else "", "text": p.text.strip()}
        for i, p in enumerate(doc.paragraphs)
        if p.text.strip()
    ]

    result = {
        "source": str(DOC_PATH),
        "image_count": len(media_filename_map),
        "product_count": len(products),
        "products": products,
        "raw_paragraph_count": len(raw_paragraphs),
    }

    OUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"Extracted {len(products)} products, {len(media_filename_map)} images")
    print(f"Written to {OUT_JSON}")


if __name__ == "__main__":
    main()
